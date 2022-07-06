import re
from Google import Create_Service
import os
import io
from googleapiclient.http import MediaIoBaseDownload
import requests
import pandas as pd
import docx
import time
import schedule
import datetime
from ast import While
import mysql.connector as con

CLIENT_SECRET_FILE = 'client_secret_Googleusercontent.json'
API_NAME = 'drive'
API_VERSION = 'v3'
SCOPES = ['https://www.googleapis.com/auth/drive']
service = Create_Service(CLIENT_SECRET_FILE, API_NAME, API_VERSION, SCOPES)

cnx = con.connect(user='root', password='', host='127.0.0.1', database='linenotify')
cur = cnx.cursor()
print(service)


def addTokenFrom_DB():
    token=[]
    cur.execute("SELECT * FROM groupline")
    records = cur.fetchall()
    for row in records:
        token.append(row[4])
    return token 



def linenotify(message, file_name=None):
    url = 'https://notify-api.line.me/api/notify'
 ##   token = ['kvLDWzoMZFWbEKdnoSgX3FX0doP1b5Id9td1fvJVnf1','CYyd1EGwn2RPoCd2zuamn93FklG8kDZPzmA0jD9A46U','mNXumpetBsDRMN7htbe9Zffj7HycrN60NtrvhuiTZRl','Uc4mZkTW3lqAMUgAelJGJSL9DLSe5YbRsuE8EcOiLG3'] # Line Notify Token CYyd1EGwn2RPoCd2zuamn93FklG8kDZPzmA0jD9A46U
    token = addTokenFrom_DB()
    data = {'message': message}
    headers = []
    for i in token :
        headers.append({'Authorization':'Bearer '+i})
    #headers = [{'Authorization':'Bearer ' + token},{'Authorization':'Bearer ' + 'CYyd1EGwn2RPoCd2zuamn93FklG8kDZPzmA0jD9A46U'}]
    session = requests.Session()
    #session_post = session.post(url, headers=headers, files=img, data =data)
    for i in headers :
        if(file_name == None):
            session_post = session.post(url, headers=i, data =data)
        else:
            #data = {'message': message,'imageFile':file_name}
            #session_post = session.post(url, headers=headers, data =data)
            img = {'imageFile': open(file_name,'rb')} #Local picture File
            session_post = session.post(url, headers=i, files=img, data =data)
        print(session_post.text)
    return

def BotLine():
        
    folder_id  = '1sxH14tKANroiMNqLCjkPJX1hoWIg5Gdc'
    query = f"parents = '{folder_id}'"


    response = service.files().list(q=query).execute()
    files = response.get('files')
    nextPageToken = response.get('nextPageToken')
    print('---------------------------')
    print(nextPageToken)
    print('----------------------------')
    while nextPageToken:
        response = service.files().list(q=query).execute()
        files.extend(response.get('files'))
        nextPageToken = response.get('nextPageToken')
        print(files,nextPageToken)
    
    file_ids = []
    file_names = []

    df = pd.DataFrame(files)
    #df = df[(df.mimeType == 'image/png') | (df.mimeType == 'image/jpg') | (df.mimeType == 'image/jpeg')]
    #df = df[df.mimeType != 'application/vnd.google-apps.document']
    df = df.sort_values(by=['name'])
    #print(df)

    file_length = len(df)
    file_ids = df.id.values.tolist()
    file_names = df.name.values.tolist()
    file_exts = df.mimeType.values.tolist()

    line_msg = ''

    for file_ids, file_names,file_exts in zip(file_ids, file_names,file_exts):
        
        request = service.files().get_media(fileId=file_ids)
        path_file = 'download/'+file_names
        #print(file_exts)
        
        if file_exts == 'application/vnd.google-apps.document':
            print("1::"+file_names)

            byteData = service.files().export_media(
                fileId=file_ids,
                mimeType = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            ).execute()
            
            with open(os.path.join(path_file+".docx"),'wb') as fw:
                fw.write(byteData)
                fw.close()
                doc = docx.Document(path_file+".docx")
                all_paras = doc.paragraphs
                len(all_paras)
                for para in all_paras:
                    line_msg += ('\n'+para.text)
                    #print(para.text)
                    #print("-------")
                
            #print(line_msg)
            if(line_msg != ""):
                time.sleep(2)
                linenotify(message=line_msg)
            
            if os.path.exists(path_file+".docx"):
                os.remove(path_file+".docx")
        
        if file_exts == 'application/vnd.google-apps.spreadsheet':
            print("2::"+file_names)
        

        if file_exts != 'application/vnd.google-apps.document' and file_exts != 'application/vnd.google-apps.spreadsheet':
            print("3::"+file_names)
            fh = io.BytesIO()
            downloader = MediaIoBaseDownload(fd=fh, request=request)
            done = False
            while not done:
                status, done = downloader.next_chunk()
                print('Download Progress {0}'.format(status.progress()*100))
            fh.seek(0)

            with open(os.path.join(path_file),'wb') as fw:
                fw.write(fh.read())            
                fw.close()
        
            if file_names.endswith('.xlsx'):
                if os.path.exists(path_file):
                    os.remove(path_file)
            
            URL = "https://docs.google.com/uc?id="
            if file_names.endswith('.PNG') or file_names.endswith('.png') or file_names.endswith('.JPG') or file_names.endswith('.jpg'):
                #notifyPicture(URL+file_ids,path_file)
                #print("IMG")
                linenotify(" ", path_file)
            if file_names.endswith('.docx'):
                doc = docx.Document(path_file+".docx")  
                all_paras = doc.paragraphs
                len(all_paras) 
                for para in all_paras:
                    line_msg += ('\n'+para.text)
                linenotify(message=line_msg)
            line_msg=""


        


        if os.path.exists(path_file):
            os.remove(path_file)
          


if __name__ == '__main__':
    # schedule.every().day.at("08:32").do(BotLine)
    BotLine()
    # while True:
    #     NowTime = datetime.datetime.now()
    #     print('While Loop', NowTime)
    #     schedule.run_pending()
    #     time.sleep(10)
    